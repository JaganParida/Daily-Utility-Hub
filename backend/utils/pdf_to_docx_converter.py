"""
PDF to DOCX Converter - Fully Editable & Exact Visual Match
Primary: pdf2docx (Industry standard for exact layout, tables, logos, images, borders, and editable text)
Fallback: Custom PyMuPDF extractor with embedded images and tables
"""
import sys, os, io, traceback

def convert_with_pdf2docx(pdf_path, docx_path):
    """
    Uses pdf2docx to recreate exact PDF layout, images, tables, shapes,
    and text formatting directly into an editable Word document.
    """
    from pdf2docx import Converter
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None, multi_processing=False)
    cv.close()

def custom_convert(pdf_path, docx_path):
    """
    Fallback extractor using PyMuPDF and python-docx.
    Extracts text, tables, and images individually.
    """
    import fitz
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pdf = fitz.open(pdf_path)
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(1.27)
        s.bottom_margin = Cm(1.27)
        s.left_margin = Cm(1.9)
        s.right_margin = Cm(1.27)

    for pi in range(len(pdf)):
        page = pdf[pi]
        pw = page.rect.width

        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        blocks.sort(key=lambda b: (round(b["bbox"][1] / 3) * 3, b["bbox"][0]))

        # Detect tables
        table_bi = set()
        tables_data = []
        try:
            found = page.find_tables()
            if found and found.tables:
                for t in found.tables:
                    tb = t.bbox
                    rows = t.extract()
                    if rows and len(rows) > 0:
                        tables_data.append({"bbox": tb, "rows": rows, "used": False})
                        for bi, b in enumerate(blocks):
                            if b["type"] == 0:
                                bb = b["bbox"]
                                if bb[0] >= tb[0]-5 and bb[1] >= tb[1]-5 and bb[2] <= tb[2]+5 and bb[3] <= tb[3]+5:
                                    table_bi.add(bi)
        except:
            pass

        # Track images
        embedded_xrefs = set()

        for bi, block in enumerate(blocks):
            if bi in table_bi:
                for td in tables_data:
                    if td["used"]:
                        continue
                    tb = td["bbox"]
                    if block["bbox"][1] >= tb[1]-5 and block["bbox"][1] <= tb[3]+5:
                        _add_table(doc, td["rows"])
                        td["used"] = True
                        break
                continue

            if block["type"] == 0:
                _add_text(doc, block, pw)

            elif block["type"] == 1:
                bbox = block["bbox"]
                w_pt = bbox[2] - bbox[0]
                h_pt = bbox[3] - bbox[1]
                if w_pt < 15 or h_pt < 15:
                    continue

                img_data = block.get("image")
                if not img_data:
                    for ii in page.get_images(full=True):
                        xref = ii[0]
                        if xref in embedded_xrefs:
                            continue
                        try:
                            rects = page.get_image_rects(xref)
                            for r in rects:
                                if abs(r.x0 - bbox[0]) < 20 and abs(r.y0 - bbox[1]) < 20:
                                    base = pdf.extract_image(xref)
                                    img_data = base.get("image")
                                    embedded_xrefs.add(xref)
                                    break
                        except:
                            pass
                        if img_data:
                            break

                if img_data:
                    try:
                        stream = io.BytesIO(img_data)
                        max_w = Inches(5.8)
                        target_w = min(Inches(w_pt / 72.0), max_w)
                        doc.add_picture(stream, width=target_w)
                    except Exception as ie:
                        print(f"Failed to add image: {ie}", file=sys.stderr)

        if pi < len(pdf) - 1:
            doc.add_page_break()


def _add_text(doc, block, pw):
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = block.get("lines", [])
    if not lines:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15

    # Check center alignment
    bx0, by0, bx1, by1 = block["bbox"]
    block_mid = (bx0 + bx1) / 2.0
    page_mid = pw / 2.0
    if abs(block_mid - page_mid) < 30 and (bx1 - bx0) < pw * 0.75:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in lines:
        for span in line.get("spans", []):
            text = span.get("text", "")
            if not text:
                continue
            run = p.add_run(text)
            flags = span.get("flags", 0)
            font_name = (span.get("font", "")).lower()
            size = span.get("size", 10.5)

            run.font.size = Pt(max(8, min(36, size)))
            if (flags & 2 != 0) or "bold" in font_name or "black" in font_name:
                run.bold = True
            if (flags & 1 != 0) or "italic" in font_name or "oblique" in font_name:
                run.italic = True

            color_int = span.get("color", 0)
            if color_int and color_int != 0:
                r = (color_int >> 16) & 0xFF
                g = (color_int >> 8) & 0xFF
                b = color_int & 0xFF
                if (r, g, b) != (0, 0, 0) and (r, g, b) != (255, 255, 255):
                    run.font.color.rgb = RGBColor(r, g, b)


def _add_table(doc, rows):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not rows or len(rows) == 0:
        return
    num_cols = max(len(r) for r in rows)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for ri, row in enumerate(rows):
        for ci in range(num_cols):
            cell = table.cell(ri, ci)
            txt = str(row[ci]).strip() if ci < len(row) and row[ci] else ""
            cell.text = txt
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if ri == 0:
                        run.bold = True


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: pdf_to_docx_converter.py <pdf> <docx>", file=sys.stderr)
        sys.exit(1)

    pdf_in, docx_out = sys.argv[1], sys.argv[2]
    if not os.path.exists(pdf_in):
        print(f"ERROR: Not found: {pdf_in}", file=sys.stderr)
        sys.exit(1)

    ok = False

    # PRIMARY: pdf2docx (Recreates exact PDF layout, logos, images, tables, shapes, and editable text)
    try:
        print("[PDF-to-Word] Converting with pdf2docx (exact layout engine)...", file=sys.stderr)
        convert_with_pdf2docx(pdf_in, docx_out)
        if os.path.exists(docx_out) and os.path.getsize(docx_out) > 200:
            ok = True
            print("[PDF-to-Word] pdf2docx conversion successful!", file=sys.stderr)
    except Exception as e:
        print(f"[PDF-to-Word] pdf2docx encountered error: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)

    # FALLBACK: Custom PyMuPDF
    if not ok:
        try:
            print("[PDF-to-Word] Fallback: Converting with PyMuPDF extractor...", file=sys.stderr)
            custom_convert(pdf_in, docx_out)
            if os.path.exists(docx_out) and os.path.getsize(docx_out) > 200:
                ok = True
                print("[PDF-to-Word] PyMuPDF conversion successful!", file=sys.stderr)
        except Exception as e:
            print(f"[PDF-to-Word] PyMuPDF failed: {e}", file=sys.stderr)
            traceback.print_exc(file=sys.stderr)

    if ok:
        print("SUCCESS")
    else:
        print("ERROR", file=sys.stderr)
        sys.exit(1)
